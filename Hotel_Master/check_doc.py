# -*- coding: utf-8 -*-
from docx import Document

doc = Document(r'c:\Users\pc\Desktop\Antigravity_Projeler\hotel-master\StudentStyle_SRS.docx')

print('=== SAYFA/BOLUM ANALIZI ===')
print(f'Toplam paragraf: {len(doc.paragraphs)}')
print(f'Toplam tablo: {len(doc.tables)}')

img_count = 0
for rel in doc.part.rels.values():
    if 'image' in rel.reltype:
        img_count += 1
print(f'Gomulu gorsel: {img_count}')

print()
print('=== BASLIKLARIN LISTESI ===')
for p in doc.paragraphs:
    if p.style.name.startswith('Heading'):
        print(f'  [{p.style.name}] {p.text[:80]}')

print()
print('=== TURKCE KARAKTER KONTROLU ===')
ascii_hatali = ['saglamalidir', 'gostermelidir', 'olmalidir', 'yapilmalidir', 'kullanilmis', 'olustur', 'icerik']
hata_sayisi = 0
for p in doc.paragraphs:
    for kelime in ascii_hatali:
        if kelime in p.text.lower():
            hata_sayisi += 1
            if hata_sayisi <= 5:
                print(f'  HATA paragraf: "{kelime}" -> ...{p.text[:70]}...')

for t in doc.tables:
    for row in t.rows:
        for cell in row.cells:
            for kelime in ascii_hatali:
                if kelime in cell.text.lower():
                    hata_sayisi += 1
                    if hata_sayisi <= 5:
                        print(f'  HATA tablo: "{kelime}" -> ...{cell.text[:60]}...')

if hata_sayisi == 0:
    print('  TEMIZ - ASCII Turkce hatasi bulunamadi!')
else:
    print(f'  TOPLAM {hata_sayisi} ASCII Turkce hatasi bulundu.')

print()
print('=== KAPAK BILGILERI ===')
kapak_kontrol = ['Ali Rıza', 'Akil Rahman', 'BUKET DOĞAN', 'BLM3008']
for arama in kapak_kontrol:
    bulundu = False
    for t in doc.tables:
        for row in t.rows:
            for cell in row.cells:
                if arama.lower() in cell.text.lower():
                    bulundu = True
    print(f'  {arama}: {"OK" if bulundu else "EKSIK!"}')

print()
print('=== GOMULU GORSELLER ===')
for rel in doc.part.rels.values():
    if 'image' in rel.reltype:
        print(f'  {rel.target_ref}')

print()
print('=== BOS TABLO KONTROLU ===')
bos = 0
for i, t in enumerate(doc.tables):
    for row in t.rows:
        for cell in row.cells:
            if cell.text.strip() == '' and len(t.rows) > 2:
                bos += 1
if bos == 0:
    print('  Bos hucre sorunu yok.')
else:
    print(f'  {bos} bos hucre bulundu (onay tablosu haric normal).')
